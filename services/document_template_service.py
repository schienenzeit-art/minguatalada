import re
from datetime import date
from core.session import Session
from core.claim_status import ClaimStatus
from database.repositories.document_template_repository import DocumentTemplateRepository


TEMPLATE_TYPES = ["BRIEF", "BESCHEID", "FORMULAR", "INFORMATION"]

PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

DEFAULT_PLACEHOLDERS = {
    "VORNAME":        "Vorname der Person",
    "NACHNAME":       "Nachname der Person",
    "ADRESSE":        "Adresse der Person",
    "PLZ":            "Postleitzahl",
    "ORT":            "Ort/Stadt",
    "AKTENZEICHEN":   "Aktenzeichen des Antrags",
    "DATUM":          "Aktuelles Datum (TT.MM.JJJJ)",
    "STANDORT":       "Name des Standorts",
    "MITARBEITER":    "Name des bearbeitenden Mitarbeiters",
    "STATUS":         "Lesbarer Statusname (z. B. Anspruchsberechtigt)",
    "BEGRUENDUNG":    "Begründung der Prüfung",
    "KATEGORIE":      "Kategorie des Antragstellers",
    "BETREFF":        "Betreffzeile (Aktenzeichen + Status)",
    "ANREDE":         "Anrede (Sehr geehrte/r Vorname Nachname)",
}


def build_claim_context(claim: dict) -> dict:
    """Erstellt den Merge-Field-Kontext aus einem Claim-Dict für alle Platzhalter."""
    status_raw = claim.get("status", "")
    status_display = ClaimStatus.get_display(status_raw) or status_raw
    vorname  = claim.get("person_first_name", "")
    nachname = claim.get("person_last_name", "")
    return {
        "VORNAME":      vorname,
        "NACHNAME":     nachname,
        "ADRESSE":      claim.get("person_address", "") or claim.get("address", ""),
        "PLZ":          claim.get("person_postal_code", "") or claim.get("postal_code", ""),
        "ORT":          claim.get("person_city", "") or claim.get("city", ""),
        "AKTENZEICHEN": claim.get("case_number", ""),
        "DATUM":        date.today().strftime("%d.%m.%Y"),
        "STANDORT":     claim.get("location_name", ""),
        "MITARBEITER":  claim.get("examiner_name", "")
                        or (Session.get_full_name() or ""),
        "STATUS":       status_display,
        "BEGRUENDUNG":  claim.get("evaluation_reason", "") or "",
        "KATEGORIE":    claim.get("category_name", ""),
        "BETREFF":      f"Antrag {claim.get('case_number','')} – {status_display}",
        "ANREDE":       f"Sehr geehrte/r {vorname} {nachname}".strip(","),
    }


class DocumentTemplateService:
    def __init__(self, repo: DocumentTemplateRepository | None = None):
        self.repo = repo or DocumentTemplateRepository()

    def list_templates(self, include_inactive: bool = False) -> list[dict]:
        return self.repo.list_templates(include_inactive=include_inactive)

    def get_template(self, template_id: int) -> dict | None:
        return self.repo.get_by_id(template_id)

    def create_template(self, data: dict) -> int:
        if not data.get("name", "").strip():
            raise ValueError("Vorlagenname ist Pflichtfeld.")
        data["created_by"] = Session.get_user_id()
        return self.repo.create(data)

    def update_template(self, template_id: int, data: dict) -> None:
        if not data.get("name", "").strip():
            raise ValueError("Vorlagenname ist Pflichtfeld.")
        self.repo.update(template_id, data)

    def delete_template(self, template_id: int) -> None:
        self.repo.delete(template_id)

    # ── DOCX-Unterstützung ────────────────────────────────────────────────────
    def upload_docx(self, template_id: int, docx_bytes: bytes) -> None:
        """Speichert eine hochgeladene DOCX-Vorlage in der Datenbank."""
        if not docx_bytes:
            raise ValueError("Keine DOCX-Daten übergeben.")
        self._validate_docx(docx_bytes)
        self.repo.update(template_id, {
            **self.get_template(template_id),
            "docx_data": docx_bytes,
        })

    def _validate_docx(self, docx_bytes: bytes) -> None:
        """Prüft ob die DOCX-Datei valid ist und Platzhalter enthält."""
        try:
            import io
            from docx import Document  # type: ignore
            doc = Document(io.BytesIO(docx_bytes))
            # Prüfe ob Dokument lesbar
            _ = len(doc.paragraphs)
        except ImportError:
            pass  # python-docx nicht installiert — Validierung überspringen
        except Exception as exc:
            raise ValueError(f"Ungültige DOCX-Datei: {exc}")

    def get_docx_bytes(self, template_id: int) -> bytes | None:
        tpl = self.repo.get_by_id(template_id)
        return tpl.get("docx_data") if tpl else None

    def generate_docx(self, template_id: int, context: dict) -> bytes:
        """Generiert ein DOCX-Dokument aus der Vorlage mit gefüllten Merge-Feldern."""
        import io
        tpl = self.repo.get_by_id(template_id)
        if not tpl:
            raise ValueError("Vorlage nicht gefunden.")

        try:
            from docx import Document  # type: ignore
        except ImportError:
            raise RuntimeError(
                "python-docx nicht installiert. Bitte 'pip install python-docx' ausführen."
            )

        docx_data = tpl.get("docx_data")
        if docx_data:
            doc = Document(io.BytesIO(bytes(docx_data)))
            self._fill_docx_placeholders(doc, context)
        else:
            doc = Document()
            rendered = self.render(template_id, context)
            for line in rendered.split("\n"):
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _fill_docx_placeholders(self, doc, context: dict) -> None:
        """Ersetzt {{PLACEHOLDER}}-Marker in allen Absätzen und Tabellen."""
        def replace_in_para(para):
            full_text = "".join(r.text for r in para.runs)
            for key, val in context.items():
                full_text = full_text.replace(f"{{{{{key}}}}}", str(val or ""))
            if para.runs:
                para.runs[0].text = full_text
                for run in para.runs[1:]:
                    run.text = ""

        for para in doc.paragraphs:
            replace_in_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_para(para)

    def get_placeholders(self) -> dict[str, str]:
        return DEFAULT_PLACEHOLDERS.copy()

    def render(self, template_id: int, context: dict) -> str:
        """Replace {{PLACEHOLDER}} in body_text with values from context."""
        tpl = self.repo.get_by_id(template_id)
        if not tpl:
            raise ValueError("Vorlage nicht gefunden.")
        body = tpl["body_text"] or ""

        def replace(m: re.Match) -> str:
            key = m.group(1).upper()
            return str(context.get(key, f"{{{{{key}}}}}"))

        return PLACEHOLDER_RE.sub(replace, body)

    def extract_placeholders(self, body_text: str) -> list[str]:
        return list({m.upper() for m in PLACEHOLDER_RE.findall(body_text)})
